# src/loaders/document_loader.py

import os
import sys
from datetime import datetime

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "../.."))
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))


# ==================================================
# CONFIG
# ==================================================

DATA_FOLDER   = "data"
OUTPUT_FOLDER = "outputs/chunks"
SUPPORTED_EXT = [".pdf", ".txt", ".docx"]


# ==================================================
# DOCUMENT LOADER CLASS
# ==================================================

class DocumentLoader:

    def __init__(self):

        print("\nInitializing Document Loader...")
        print(f"  Data folder   : {DATA_FOLDER}")
        print(f"  Output folder : {OUTPUT_FOLDER}")
        print("Document Loader ready!\n")


    # ==================================================
    # LOAD SINGLE PDF
    # ==================================================

    def load_pdf(self, pdf_path: str) -> str:

        try:
            from pypdf import PdfReader
            import concurrent.futures

            reader = PdfReader(pdf_path)

            def extract_page(page):
                try:
                    return page.extract_text() or ""
                except Exception:
                    return ""

            with concurrent.futures.ThreadPoolExecutor(
                max_workers = os.cpu_count()
            ) as executor:
                pages = list(
                    executor.map(extract_page, reader.pages)
                )

            text       = "\n".join(pages)
            page_count = len(reader.pages)

            print(
                f"  Loaded PDF : {os.path.basename(pdf_path)}"
                f" ({page_count} pages, {len(text)} chars)"
            )

            return text

        except Exception as e:
            print(f"  PDF error: {e}")
            return ""


    # ==================================================
    # LOAD SINGLE TXT
    # ==================================================

    def load_txt(self, txt_path: str) -> str:

        try:
            with open(txt_path, "r", encoding="utf-8") as f:
                text = f.read()

            print(
                f"  Loaded TXT : {os.path.basename(txt_path)}"
                f" ({len(text)} chars)"
            )

            return text

        except Exception as e:
            print(f"  TXT error: {e}")
            return ""


    # ==================================================
    # LOAD SINGLE DOCX
    # ==================================================

    def load_docx(self, docx_path: str) -> str:

        try:
            import docx

            doc   = docx.Document(docx_path)
            lines = []

            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    ])
                    if row_text:
                        lines.append(row_text)

            text = "\n".join(lines)

            print(
                f"  Loaded DOCX: {os.path.basename(docx_path)}"
                f" ({len(text)} chars)"
            )

            return text

        except Exception as e:
            print(f"  DOCX error: {e}")
            return ""


    # ==================================================
    # LOAD SINGLE DOCUMENT
    # ==================================================

    def load_document(self, file_path: str) -> dict:

        ext  = os.path.splitext(file_path)[1].lower()
        name = os.path.basename(file_path)

        print(f"\nLoading: {name}")

        if ext == ".pdf":
            text = self.load_pdf(file_path)

        elif ext == ".txt":
            text = self.load_txt(file_path)

        elif ext == ".docx":
            text = self.load_docx(file_path)

        else:
            print(f"  Unsupported: {ext}")
            return {}

        if not text.strip():
            print(f"  Empty document: {name}")
            return {}

        return {
            "file_name"  : name,
            "file_path"  : file_path,
            "file_type"  : ext.replace(".", ""),
            "text"       : text,
            "char_count" : len(text),
            "word_count" : len(text.split()),
            "loaded_at"  : datetime.now().isoformat()
        }


    # ==================================================
    # LOAD ALL DOCUMENTS FROM FOLDER
    # ==================================================

    def load_folder(
        self,
        folder_path : str = DATA_FOLDER
    ) -> list:

        if not os.path.exists(folder_path):
            print(f"  Folder not found: {folder_path}")
            return []

        all_files = []

        for file in sorted(os.listdir(folder_path)):

            ext = os.path.splitext(file)[1].lower()

            if ext in SUPPORTED_EXT:
                all_files.append(
                    os.path.join(folder_path, file)
                )

        if not all_files:
            print(f"  No documents in {folder_path}")
            return []

        print(f"\nFound {len(all_files)} documents")
        print("=" * 60)

        documents = []

        for i, file_path in enumerate(all_files):

            print(f"\n[{i+1}/{len(all_files)}]", end=" ")

            doc = self.load_document(file_path)

            if doc:
                documents.append(doc)

        print(f"\n{'='*60}")
        print(f"Loaded: {len(documents)}/{len(all_files)}")

        return documents


    # ==================================================
    # VALIDATE DOCUMENTS
    # ==================================================

    def validate_documents(self, documents: list) -> list:

        if not documents:
            return []

        valid   = []
        invalid = []

        for doc in documents:

            issues = []

            if doc["char_count"] < 100:
                issues.append("Too short")

            if doc["word_count"] < 20:
                issues.append("Too few words")

            if issues:
                invalid.append({
                    "file"   : doc["file_name"],
                    "issues" : issues
                })
            else:
                valid.append(doc)

        print(f"\n  Validation:")
        print(f"  Valid   : {len(valid)}")
        print(f"  Invalid : {len(invalid)}")

        if invalid:
            for inv in invalid:
                print(
                    f"  {inv['file']}: "
                    f"{', '.join(inv['issues'])}"
                )

        # If all invalid return all anyway
        if not valid and documents:
            print("  Using all documents despite warnings")
            return documents

        return valid


    # ==================================================
    # GET STATS
    # ==================================================

    def get_stats(self, documents: list) -> dict:

        if not documents:
            return {
                "total_documents" : 0,
                "total_chars"     : 0,
                "total_words"     : 0,
                "by_type"         : {},
                "avg_chars"       : 0,
                "avg_words"       : 0,
            }

        total_chars = sum(d["char_count"] for d in documents)
        total_words = sum(d["word_count"] for d in documents)

        by_type = {}
        for doc in documents:
            ft = doc["file_type"]
            if ft not in by_type:
                by_type[ft] = 0
            by_type[ft] += 1

        return {
            "total_documents" : len(documents),
            "total_chars"     : total_chars,
            "total_words"     : total_words,
            "by_type"         : by_type,
            "avg_chars"       : total_chars // len(documents),
            "avg_words"       : total_words // len(documents),
        }


    # ==================================================
    # PRINT STATS
    # ==================================================

    def print_stats(self, stats: dict):

        if not stats or stats.get("total_documents", 0) == 0:
            print("  No stats available")
            return

        print(f"\n{'='*60}")
        print(f"  DOCUMENT LOADER STATS")
        print(f"{'='*60}")
        print(f"  Total Documents : {stats['total_documents']}")
        print(f"  Total Chars     : {stats['total_chars']:,}")
        print(f"  Total Words     : {stats['total_words']:,}")
        print(f"  Avg Chars/Doc   : {stats['avg_chars']:,}")
        print(f"  Avg Words/Doc   : {stats['avg_words']:,}")
        print(f"\n  By File Type:")

        for ftype, count in stats["by_type"].items():
            print(f"     .{ftype} : {count} files")

        print(f"{'='*60}\n")


    # ==================================================
    # SAVE DOCUMENTS
    # ==================================================

    def save_documents(
        self,
        documents   : list,
        output_path : str = OUTPUT_FOLDER
    ):

        os.makedirs(output_path, exist_ok=True)

        saved = 0

        for doc in documents:

            filename = (
                doc["file_name"]
                .replace(".pdf",  "_raw.txt")
                .replace(".docx", "_raw.txt")
                .replace(".txt",  "_raw.txt")
            )

            save_path = os.path.join(output_path, filename)

            try:
                with open(
                    save_path, "w", encoding="utf-8"
                ) as f:
                    f.write(f"FILE   : {doc['file_name']}\n")
                    f.write(f"TYPE   : {doc['file_type']}\n")
                    f.write(f"CHARS  : {doc['char_count']}\n")
                    f.write(f"WORDS  : {doc['word_count']}\n")
                    f.write(f"LOADED : {doc['loaded_at']}\n")
                    f.write("=" * 80 + "\n\n")
                    f.write(doc["text"])

                saved += 1
                print(f"  Saved → {save_path}")

            except Exception as e:
                print(f"  Save error: {e}")

        print(f"\nSaved {saved}/{len(documents)} documents")


# ==================================================
# MAIN
# ==================================================

if __name__ == "__main__":

    loader = DocumentLoader()

    # Load all documents
    documents = loader.load_folder(DATA_FOLDER)

    if not documents:
        print("No documents loaded!")
        exit()

    # Validate
    documents = loader.validate_documents(documents)

    # Stats
    stats = loader.get_stats(documents)
    loader.print_stats(stats)

    # Preview first doc
    if documents:
        print(f"\nFirst Document Preview:")
        print(f"{'─'*60}")
        print(f"Name  : {documents[0]['file_name']}")
        print(f"Words : {documents[0]['word_count']}")
        print(f"Text  : {documents[0]['text'][:300]}...")
        print(f"{'─'*60}")